from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E8E8E8')
    p._p.pPr.append(shading) if p._p.pPr is not None else None
    return p

def divider(doc):
    doc.add_paragraph('─' * 72)

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('OxySync Device — Encryption Setup Instructions')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x35, 0x5E)

doc.add_paragraph()

# Meta
meta = [
    ('For:', 'Fred Faiz'),
    ('From:', 'Bharat'),
    ('Date:', 'July 31, 2026'),
    ('Device:', 'ESP32 DevKit V1 (NEW TEST UNIT ONLY)'),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.add_run(label + ' ').bold = True
    r1 = p.runs[0]
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

divider(doc)

# ── Warning ──
p = doc.add_paragraph()
run = p.add_run('⚠️  CRITICAL WARNING — READ BEFORE STARTING')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

warnings = [
    'This process permanently modifies the device hardware (eFuse registers).',
    'It CANNOT be undone — ever.',
    'Do this ONLY on the new test unit — NOT on any unit currently at the client site.',
    'If the process is interrupted midway, the device may become permanently unusable.',
    'Do NOT disconnect the USB cable during any step.',
]
for w in warnings:
    p = doc.add_paragraph(w, style='List Bullet')
    p.runs[0].font.size = Pt(11)

divider(doc)

# ── Files from Bharat ──
heading(doc, 'WHAT YOU RECEIVE FROM BHARAT', level=2)
para(doc, 'Create a folder on your Desktop called OxySync_Encryption and place all these files in it:')
files = [
    'oxysync_production.ino.bootloader.signed.bin',
    'oxysync_production.ino.signed.bin',
    'oxysync_production.ino.partitions.bin',
    'boot_app0.bin',
    'oxysync_public_key.pem',
]
for f in files:
    p = doc.add_paragraph(f, style='List Bullet')
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)

# ── Steps ──
steps = [
    {
        'title': 'STEP 1 — Install Python (skip if already installed)',
        'content': [
            ('text', 'Go to https://www.python.org/downloads/ and download the latest Python for Windows.'),
            ('text', 'Important: During installation, tick the box "Add Python to PATH".'),
            ('text', 'Open Command Prompt and verify:'),
            ('code', 'python --version'),
            ('text', 'You should see a version number like Python 3.x.x'),
        ]
    },
    {
        'title': 'STEP 2 — Install esptool',
        'content': [
            ('text', 'Open Command Prompt and run:'),
            ('code', 'pip install esptool'),
            ('text', 'Wait for it to finish. Then verify:'),
            ('code', 'python -m esptool version'),
            ('text', 'You should see esptool.py v5.x.x or similar.'),
        ]
    },
    {
        'title': 'STEP 3 — Connect the device and find the COM port',
        'content': [
            ('text', '1. Connect the ESP32 test unit to your PC via USB.'),
            ('text', '2. Open Device Manager (right-click Start → Device Manager).'),
            ('text', '3. Look under Ports (COM & LPT).'),
            ('text', '4. Note the COM port number — e.g. COM3 or COM6.'),
            ('text', 'You will replace COMX in all commands below with your actual COM port number.'),
        ]
    },
    {
        'title': 'STEP 4 — Navigate to the files folder',
        'content': [
            ('text', 'Open Command Prompt and run:'),
            ('code', 'cd C:\\Users\\%USERNAME%\\Desktop\\OxySync_Encryption'),
        ]
    },
    {
        'title': 'STEP 5 — Burn the Secure Boot public key into eFuse  ⚠️ PERMANENT',
        'content': [
            ('text', 'Run the command below. Replace COMX with your COM port number.'),
            ('code', 'python -m espefuse --port COMX burn_key BLOCK2 oxysync_public_key.pem SECURE_BOOT_V2_RSA_3072'),
            ('text', 'The tool will ask you to type BURN to confirm. Type it exactly and press Enter.'),
        ]
    },
    {
        'title': 'STEP 6 — Enable Secure Boot V2  ⚠️ PERMANENT',
        'content': [
            ('text', 'Run:'),
            ('code', 'python -m espefuse --port COMX burn_efuse ABS_DONE_1'),
            ('text', 'Type BURN to confirm when prompted.'),
        ]
    },
    {
        'title': 'STEP 7 — Flash the signed firmware',
        'content': [
            ('text', 'Run this single command (copy and paste the entire line):'),
            ('code', 'python -m esptool --port COMX --baud 460800 write_flash 0x1000 oxysync_production.ino.bootloader.signed.bin 0x8000 oxysync_production.ino.partitions.bin 0xe000 boot_app0.bin 0x10000 oxysync_production.ino.signed.bin'),
            ('text', 'Wait for it to complete. You should see: Hash of data verified. / Hard resetting via RTS pin...'),
        ]
    },
    {
        'title': 'STEP 8 — Enable Flash Encryption  ⚠️ PERMANENT',
        'content': [
            ('text', 'The device will auto-encrypt its flash on next boot.'),
            ('code', 'python -m espefuse --port COMX burn_efuse FLASH_CRYPT_CNT'),
            ('text', 'Type BURN to confirm.'),
        ]
    },
    {
        'title': 'STEP 9 — Reboot and verify',
        'content': [
            ('text', '1. Unplug and replug the USB cable (or press the EN/Reset button).'),
            ('text', '2. The device will take slightly longer than usual on first boot — this is normal (it is encrypting the flash).'),
            ('text', '3. On your phone or laptop, connect to WiFi: O2_Controller  (password: techonly123)'),
            ('text', '4. Open browser and go to: 192.168.4.1'),
            ('text', '5. Login: username tech / password oxygen'),
            ('text', '6. Confirm the dashboard loads and sensor readings appear.'),
        ]
    },
    {
        'title': 'STEP 10 — Report back to Bharat',
        'content': [
            ('text', 'Send Bharat a WhatsApp message confirming:'),
            ('text', '☐  Dashboard loaded successfully'),
            ('text', '☐  Sensor reading is showing'),
            ('text', '☐  Device WiFi is working'),
            ('text', '☐  Any errors or unexpected behaviour (if any)'),
        ]
    },
]

for step in steps:
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(step['title'])
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x35, 0x5E)

    for item_type, item_text in step['content']:
        if item_type == 'code':
            code_block(doc, item_text)
        else:
            p = doc.add_paragraph(item_text)
            p.runs[0].font.size = Pt(11)

divider(doc)

# ── Future updates note ──
p = doc.add_paragraph()
run = p.add_run('IMPORTANT — Future Firmware Updates')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

para(doc, 'From this point on, only signed firmware from Bharat can be flashed to this device. Any unsigned firmware will be rejected and the device will not boot. Always contact Bharat for any future firmware updates.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('OxySync Encryption Setup — Bharat — July 31, 2026')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'C:\Bharat\Oxygen Projects\Oxy_Sync\OxySync_Encryption_Instructions_Fred.docx')
print("Done! File saved.")
